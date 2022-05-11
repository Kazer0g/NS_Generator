# =============================================import)
import os

import random

from docx import Document
from docx.shared import Inches


from tkinter import *
import tkinter as tk

# ----------------------------------------К десятичной)
def NS_to_int (n, radix):
	return int(n, radix)
# ---------------------------------------Из десятичной)
def int_to_NS (n,radix):
	digits="0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ"
	r=""
	while(n>0):
		k=n%radix   # очередная цифра
		r=digits[k]+r # приклеим к результату
		n=n//radix	
	return r
# ===============================================Equal)
class Equal:
	def __init__(self, act_num, left_cc, right_cc, left_num, right_num, left_mark, right_mark, list_answers):

		self.act_num = act_num

		self.left_cc = left_cc
		self.right_cc = right_cc

		self.left_full = left_num
		self.right_full = right_num

		self.left_mark = left_mark
		self.right_mark = right_mark

		self.list_answers = list_answers

# ========================================Main Window)
class Main_Window (tk.Tk):
	def __init__(self):
		super().__init__()


		self.list_NS = ()

		self.lot = IntVar()

		self.list_Equals = []

		# --------------------------------------Data)
		self.minval = ''
		self.maxval = ''

		self.difficult = ''

		self.times = ''

		self.mod = ''

		# -------------------------------------Input)
		self.btn_saveNS = tk.Button (self, text='Сахранить', width=10, height=1, command=self.saveNS)
		self.btn_saveNS.place (x=0, y=20)

		self.ent_list_strNS = tk.Entry (self, width=80)
		self.ent_list_strNS.place (x=0, y=0)

		self.lbl_minmax = tk.Label (self, text='Диапазон от до') 
		self.lbl_minmax.place (x=0, y=50)

		self.ent_minval = tk.Entry (self, width=12)
		self.ent_minval.place (x=0, y=70)
		self.ent_maxval = tk.Entry (self, width=12)
		self.ent_maxval.place (x=0, y=90)

		self.lbl_difficult = tk.Label (self, text='Сложность 0-100') 
		self.lbl_difficult.place (x=100, y=50)

		self.ent_difficult = tk.Entry (self, width=12)
		self.ent_difficult.place (x=100, y=70)

		self.lbl_times = tk.Label (self, text='Кол-во') 
		self.lbl_times.place (x=200, y=50)

		self.ent_times = tk.Entry (self, width=12)
		self.ent_times.place (x=200, y=70)

		self.lotbtn = tk.Checkbutton (self, text='Одно решение', variable=self.lot)
		self.lotbtn.place (x=0, y=150)

		self.btn_generate = tk.Button (self, text='Создать', command=self.generate)
		self.btn_generate.place (x=0, y=200)

		# ---------------------------------------Tests)
		self.btn_Test1 = tk.Button (self, text='Тест 1', width=10, height=1, command=self.Test1)
		self.btn_Test1.place (x=400, y=60)

	# -------------------------------------------Test1)
	def Test1 (self):
		self.ent_list_strNS.delete(0, tk.END)
		self.ent_list_strNS.insert(0, [2, 8, 10, 16])

		self.ent_minval.delete(0, tk.END)
		self.ent_minval.insert(0, 111)

		self.ent_maxval.delete(0, tk.END)
		self.ent_maxval.insert(0, 1111)

		self.ent_difficult.delete(0, tk.END)
		self.ent_difficult.insert(0, 50)

		self.ent_times.delete(0, tk.END)
		self.ent_times.insert(0, 1)
	
	# ---------------------Сохранение систем счисления)
	def saveNS (self):
		self.list_NS = list(set(str(self.ent_list_strNS.get()).split(' ')))
		for i in range (len(self.list_NS)):
			self.list_NS[i] = int(self.list_NS[i])
		self.list_NS.sort()
		self.ent_list_strNS.delete(0, tk.END)
		self.ent_list_strNS.insert(0, self.list_NS)
		print (self.list_NS)

	# -------------------------------------Сбор данных)
	def collect (self):
		self.minval = int(self.ent_minval.get()) # -------------Минимально
		self.maxval = int(self.ent_maxval.get()) # -----------Максимальное
		self.difficult = int(self.ent_difficult.get()) # --------Сложность
		self.times = int(self.ent_times.get()) # ----------Кол-во равенств
		self.mod = self.lot.get() # ---------------------------------Режим 

	# --------------------Создание открытого равенства)
	def randNum (self, equal):
		equal.act_num = random.randint(self.minval, self.maxval)
		equal.left_cc = self.list_NS[random.randint(0, len(self.list_NS))-1]
		equal.right_cc = self.list_NS[random.randint(0, len(self.list_NS)-1)]
		while equal.left_cc == equal.right_cc:
			equal.right_cc = self.list_NS[random.randint(0, len(self.list_NS)-1)]

		left_cc = equal.left_cc
		right_cc = equal.right_cc
		act_num = equal.act_num
		equal.left_full = int_to_NS(act_num, left_cc)
		equal.right_full = int_to_NS(act_num, right_cc)

	# ------------------------------Закрашивание числа)
	def markerNum (self, num):
		count = round((len(num)/100)*self.difficult)
		for i in range (count):
			index = random.randint(0, len(num)-1)
			while num[index]=='*':
				index = random.randint(0, len(num)-1)
			num = num[:index] + '*' + num[index+1:]
		return num
	
	# -------------------------Процентное закрашивание)
	def markNum (self, equal):
		equal.left_mark = self.markerNum(str(equal.left_full))
		equal.right_mark = self.markerNum(str(equal.right_full))
	
	# ----------------------Проверка возможного ответа)
	def check_valans (self, equal, left, right):
		correct_right = int_to_NS(NS_to_int (left, equal.left_cc), equal.right_cc)
		for i in range(len(right)):
			if (len(correct_right) == len(right)) and (right[i] != correct_right[i]) and (right[i] != '*'):
				return False
		return True
	
	#---------------------------- Все варианты решения)
	def all_answers (self, equal):
		count = 0
		for i in equal.left_mark:
			if i == '*':
				count = count + 1

		for i in range(equal.left_cc**count):
			left = equal.left_mark
			right = equal.right_mark
			ins = str(int_to_NS(i, equal.left_cc))
			for j in range(count-len(str(ins))):
				ins = '0' + ins
			for l in range(count):
				left = left.replace ('*', ins[l], 1)
			if self.check_valans (equal, left, right) == True:
				equal.list_answers.append(left)

	# ----------------------------Сокращение вариантов)	
	def cut_variants (self, equal):
		while len(equal.list_answers)>1:
			index = equal.left_mark.find('*')
			if index != -1:
				equal.left_mark = equal.left_mark.replace('*', equal.left_full[index], 1)
				print (equal.left_mark)
				print (equal.right_mark)
				print (equal.list_answers)
				equal.list_answers.clear()
				self.all_answers(equal)
			else:
				self.generate()

	# ----------------------------------Запись в файлы)
	def write (self):
		# -----------------------Создание папки тестов)
		try:
			os.mkdir('Tests')
			os.chdir('Tests')
		except FileExistsError:
			os.chdir('Tests')

		# -----------------------Создание папки ребуса)
		flag = False
		i = 1
		while flag == False:
			try:
				os.mkdir('Ребусы ' + str(i))
				flag = True
			except FileExistsError:
				i += 1
		
		os.chdir('Ребусы ' + str(i))

		# ------------------------------Запись ответов)
		answers = Document()
		answers.add_heading('Цифровая электроника', 0)
		# ------------------------------Запись заданий)
		tasks = Document()
		tasks.add_heading ('Цифровая электроника', 0)
		tasks.add_heading ('Ученика(цы) __ класса \'___\' _______________________ ', 2)

	def generate (self):
		self.saveNS()
		self.collect()

		for i in range(self.times):
			equal = Equal(0, 0, 0, '', '', '', '', [])
			self.list_Equals.append(equal)
			self.randNum(equal)
			self.markNum(equal)
			self.all_answers (equal)
			if self.mod == 1:
				self.cut_variants (equal)
			
		for i in self.list_Equals:
			print ('---------------------------------')
			print (i.act_num)
			print (str(i.left_cc) + '|' + str(i.right_cc))
			print (str(i.left_full) + '=' + str(i.right_full))
			print (i.left_mark + '=' + i.right_mark)
			print (equal.list_answers)
			
		# self.write()
		# os.chdir("..")
		# os.chdir("..")	
		self.list_Equals.clear()
# ++++++++++++++++++++++++++++++++++++++++++++++++++++) 
if __name__ == '__main__':
	Main_window = Main_Window()
	Main_window.title ('Ns_generator')
	Main_window.geometry ('500x300')

	Main_window.mainloop ()
# +++++++++++++++++++++++++++++++++++++++++++++++++++++)