import os

import random

from docx import Document
from docx.shared import Inches


from tkinter import *
import tkinter as tk


class Main_Window (tk.Tk):
	def __init__(self):
		super().__init__()

		self.actnum = 1

		self.cc = []
		self.mm = []

		self.li_full = []
		self.li_mark = []

		self.str_full = ""
		self.str_mark = ""
		self.st = ""

		self.left = ""
		self.right = ""
		self.left_cc = ""
		self.right_cc = ""

		self.variants = []

		self.li_answears = []

		#-------------------------------------------------------------------------------------------------------------------Выбор один/много ответов)
		self.lot = IntVar()

		self.lotbtn = tk.Checkbutton (self, text="Одно решение", variable=self.lot)
		self.lotbtn.place (x=0, y=100)


		#------------------------------------------------------------------------------------------------------------Выбор активных систем счисления)
		self.lbl_numeralsystem = tk.Label (self, text="Системы счисления") 
		self.lbl_numeralsystem.place (x=0, y=0)

		self.c_2 = IntVar()
		self.c_10 = IntVar()
		self.c_16 = IntVar()

		self.cbtn_2 = tk.Checkbutton (self, text="2", variable=self.c_2)
		self.cbtn_2.place (x=0, y=20)
		self.cbtn_10 = tk.Checkbutton (self, text="10", variable=self.c_10)
		self.cbtn_10.place (x=0, y=40)
		self.cbtn_16 = tk.Checkbutton (self, text="16", variable=self.c_16)
		self.cbtn_16.place (x=0, y=60)
		
		#------------------------------------------------------------------------------------------------------------------Выбор диапазона значений)
		self.lbl_minmax = tk.Label (self, text="Диапазон от до") 
		self.lbl_minmax.place (x=160, y=0)

		self.ent_minval = tk.Entry (self, width=12)
		self.ent_minval.place (x=160, y=20)
		self.ent_maxval = tk.Entry (self, width=12)
		self.ent_maxval.place (x=160, y=40)

		#---------------------------------------------------------------------------------------------------------------------------Выбор сложности)
		self.lbl_difficult = tk.Label (self, text="Сложность 0-100") 
		self.lbl_difficult.place (x=280, y=0)

		self.ent_difficult = tk.Entry (self, width=12)
		self.ent_difficult.place (x=280, y=20)

		#---------------------------------------------------------------------------------------------------------------------------Кол-во примеров)
		self.lbl_times = tk.Label (self, text="Кол-во") 
		self.lbl_times.place (x=400, y=0)


		self.ent_times = tk.Entry (self, width=12)
		self.ent_times.place (x=400, y=20)
		#---------------------------------------------------------------------------------------------------------------------------Кнопка создания)
		self.btn_generate = tk.Button (self, text="Создать", command=self.generate)
		self.btn_generate.place (x=0, y=165)

	#----------------------------------------------------------------------------------------------------------------------Функция проверки integer)
	def int_check (self, a):
		try:
			int(a)
			return True
		except ValueError:
			return False
	
	#----------------------------------------------------------------------------------------------------------Функция создания дирикторий и файлов)
	def generate (self):
		self.createNum ()
		self.createTests ()
		self.createnewdir ()
		self.createfiles ()
		
		os.chdir("..")
		os.chdir("..")

		self.li_full.clear()
		self.li_mark.clear()
	
	#---------------------------------------------------------------------------------------------------------------------Функция создания равенства)
	def createNum (self):
		self.collectCC ()
		print ("Выбранные системы: ", self.cc)
		self.collectMM()
		print ("Max/Min:", self.mm)
		for i in range (int(self.ent_times.get())):
			self.rand()
			self.mark (self.left, self.right)
			if self.lot.get() != 1:
				self.li_full.append(str(i+1)+ ") " + "Загаданное число в десятичной системе:" + str(self.actnum))
				self.minvar()
				self.li_mark.append(self.str_mark)
			elif self.lot.get() == 1:
				self.li_full.append(str(i+1)+ ") " + "Загаданное число в десятичной системе:" + str(self.actnum)+ "\n" + self.str_full)
				self.minvar ()
		print(self.li_full)
		print (self.li_mark)
		
	#------------------------------------------------------------------------------------------------------Функция сбора актуальных систем счисления)
	def collectCC (self):
		self.cc = []
		if self.c_2.get() == 1:
			self.cc.append (2)
		if self.c_10.get() == 1:
			self.cc.append(10)
		if self.c_16.get() == 1:
			self.cc.append(16)

	#--------------------------------------------------------------------------------------------------------------Функция сведения ответов к одному)
	def minvar (self):
		lr = self.str_mark.split("=")
		lrr = self.str_full.split("=")
		print (lrr)
		print (lr)
		left = lr[0]
		leftr = lr[0]
		right = lr[1]
		rightr = lr[1]
		if left[0] == "0":
			left = left[2:]
		if right[0] == "0":
			right = right[2:]

		if len(left) >= len(right):
			self.variant (right, self.right_cc, lr[0], self.left_cc, lrr[0], lr[1])
		else:
			self.variant (left, self.left_cc, lr[1], self.right_cc, lrr[1], lr[0])

		
	#-----------------------------------------------------------------------------------------------------------------------Функция поиска вариантов)
	def variant (self, num, cc_num, side_num, cc_side, check, writer):
		print (cc_num, " ", cc_side)
		print (num)
		print (writer)
		count = 0
		side_vars = []
		pops = []
		self.variants.clear()
		for i in str(num):
			if i == "*":
				count = count + 1

		for calk in range (cc_num**count):
			
			if cc_num == 2:
				cc_calk = bin(calk)
			elif cc_num == 10:
				cc_calk = int(calk)
			elif cc_num == 16:
				cc_calk = hex(calk)
			str_calk = str(cc_calk)
			str_num = str(num)

			if (len(str_calk) > 1) and (str_calk[1] == "x"):
				while len(str_calk[2:]) < count:
					str_calk = str_calk[:2] + "0" + str_calk[2:]
			else:
				while len(str_calk) < count:
					str_calk = "0" + str_calk

			if (len(str_calk) > 1) and (str_calk[1] == "x"):
				for letter in str_calk[2:]:
					str_num = str_num.replace ("*", letter, 1)
			else:
				for letter in str_calk:
					str_num = str_num.replace ("*", letter, 1)

			if (str_num[0] == "0") and (len(str_num)>2):
				str_num_base = str_num[2:]
			else:
				str_num_base = str_num

			
			if cc_side == 16:
				side_var = hex(int(str_num_base, cc_num))
			elif cc_side == 10:
				side_var = int(str_num_base, cc_num)
			elif cc_side == 2:
				side_var = bin(int(str_num_base, cc_num))


			if len(side_num) == len(str(side_var)):
				for l in range(len(side_num)):
					if (side_num[l] == str(side_var)[l]) or (side_num[l] == "*"):
						flag = True
					else:
						flag = False
						break

				if flag == True:
					self.variants.append(str_num)
					side_vars.append(side_var)
					print (side_num, str(side_var))
		
		if self.lot.get() == 1:
			self.cut(num, cc_num, side_num, cc_side, check, writer, side_vars, pops)
		else:
			lo = ""
			for u in range(len (side_vars)):
				lo = lo + self.variants[u] + "=" + str(side_vars[u]) + "\n"
			self.li_full.append(lo)

	def cut (self, num, cc_num, side_num, cc_side, check, writer, side_vars, pops):

		print (self.variants)
		print (side_vars)

		while len(self.variants) > 1:
			index = side_num[::-1].find ("*")
			if index != 0:
				side_num = side_num[:-index-1] + check[-index-1] + side_num[-index:]
			else:
				side_num = side_num[:-index-1] + check[-index-1]


			for s in range(len(side_vars)):
				for l in range(len(side_num)):
					if (side_num[l] == str(side_vars[s])[l]) or (side_num[l] == "*"):
						flag = True
					else:
						flag = False
						pops.append(s)
						break
						
			if len(pops) >= 1:
				x = 0 
				for m in pops:
					side_vars.pop(m-x)
					self.variants.pop(m-x)
					x = x + 1
			pops.clear()


			print (side_num)
			print (check)

			print (self.variants)
			print (side_vars)

		self.li_mark.append(str(writer + "=" + side_num))

	#------------------------------------------------------------------------------------------Функция сбора диапазона значений генерируегмого числа)
	def collectMM (self):
		numMin = self.ent_minval.get()
		numMax = self.ent_maxval.get()
		self.mm = []
		if self.int_check (numMin) == True:
			numMin = int(numMin)
			self.mm.append(numMin)
		else:
			print ("Некоректно введено минимальное значение")
			self.mm = []

		if self.int_check (numMax) == True:
			numMax = int(numMax)
			self.mm.append(numMax)
		else:
			print ("Некоректно введено максимальное значение")
			self.mm = []

		if numMin > numMax:
			print ("Некоректно введен диапазон")

	#---------------------------------------------------------------------------------------------------------Функция генерации рандомного равенства)
	def rand (self):
		self.actnum = random.randint(self.mm[0], self.mm[1])
		self.left_cc = self.cc[random.randint(0, len(self.cc))-1]
		self.right_cc = self.left_cc
		while self.left_cc == self.right_cc:
			self.right_cc = self.cc[random.randint(0, len(self.cc))-1]	
		if self.left_cc == 2:
			left = bin(self.actnum)
		elif self.left_cc == 10:
			left = int(self.actnum)
		elif self.left_cc ==16:
			left = hex(self.actnum)
		if self.right_cc == 2:
			right = bin(self.actnum)
		elif self.right_cc == 10:
			right = int(self.actnum)
		elif self.right_cc ==16:
			right = hex(self.actnum)

		self.left = left
		self.right = right

		self.str_full = str(left) + "=" + str(right)

	#-------------------------------------------------------------------------------------------------------Функция генерации закрашенного равенства)
	def mark (self, left, right):
		count = []
		procent = int(self.ent_difficult.get())
		self.str_mark = self.marker(left, procent) + "=" + self.marker(right, procent)

	#----------------------------------------------------------------------------------------------------------------------Функция расчета сложности)
	def marker (self, num, procent):
		self.st = str(num)

		pref = ""

		if self.st[0] == "0":
			pref = self.st[:2]
			self.st = self.st[2:]

		count = round((len(self.st)/100)*procent)
		for i in range (count):
			self.markering (self.st, procent)

		return pref + self.st
	
	#-----------------------------------------------------------------------------------------------------------------Функция закрашивания равенства)
	def markering (self, st, procent):
		index = random.randint(0, len(st)-1)
		if st[index] == "*":
			self.markering (st, procent)
		elif st[index] != "*":
			self.st = st[:index] + "*" + st[index+1:]
		return self.st

	#---------------------------------------------------------------------------------------------------------------Функция создания папки с тестами)
	def createTests (self):
		try:
			os.mkdir("Tests")
			os.chdir("Tests")
		except FileExistsError:
			os.chdir("Tests")
	
	#--------------------------------------------------------------------------------------------------------Функция создания папки с новыми тестами)
	def createnewdir (self):
		flag = False
		i = 1
		while flag == False:
			try:
				os.mkdir("Ребусы " + str(i))
				flag = True
			except FileExistsError:
				i += 1
		
		os.chdir("Ребусы " + str(i))
	
	#------------------------------------------------------------------------------------------------------------------------Функция создания файлов)
	def createfiles (self):
		#------------------------------------------------------------------------------------------------------------------------------------Задания)
		tasks = Document()
		tasks.add_heading ("Цифровая электроника", 0)
		tasks.add_heading ("Ученика(цы) __ класса \"___\" _______________________ ", 2)
		for i in range(len(self.li_mark)):
			so = str(str(i+1) + ") " + self.li_mark[i])
			tasks.add_paragraph (so)

		tasks.save("Задачи.docx")

		#------------------------------------------------------------------------------------------------------------------------------------Ответы)
		answers = Document()
		answers.add_heading("Цифровая электроника", 0)
		for j in range(len(self.li_full)):
			answers.add_paragraph(self.li_full[j])
		answers.save("Ответы.docx")
		

if __name__ == "__main__":
	main_window = Main_Window ()

	main_window.title ("NS")
	main_window.geometry ("500x400")

	main_window.mainloop()